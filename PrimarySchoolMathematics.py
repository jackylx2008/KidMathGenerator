import random
import yaml
import os
import logging
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from logging_config import setup_logger
import convert_to_pdf


class MathQuizGenerator:
    def __init__(self, config_path="config.yaml"):
        # 加载配置
        with open(config_path, "r", encoding="utf-8") as f:
            self.config = yaml.safe_load(f)

        # 初始化日志
        log_level_str = self.config.get("log_level", "INFO").upper()
        log_level = getattr(logging, log_level_str, logging.INFO)
        self.logger = setup_logger(log_level=log_level)
        self.logger.info("初始化小学口算题生成器...")

    def generate_problem(self):
        """根据配置随机生成一道题目"""
        settings = self.config["quiz"]["settings"]
        setting = random.choice(settings)

        steps = setting.get("steps", 1)
        t1_min = setting.get("term1_min", 1)
        t1_max = setting.get("term1_max", 100)
        t2_min = setting.get("term2_min", 1)
        t2_max = setting.get("term2_max", 100)
        # 支持 term3 配置
        t3_min = setting.get("term3_min", 1)
        t3_max = setting.get("term3_max", 100)

        # 兼容旧版的 operators，支持新版的 operators1, operators2
        ops_pool = setting.get("operators", ["+"])
        ops_pool1 = setting.get("operators1", ops_pool)
        ops_pool2 = setting.get("operators2", ops_pool)

        r_min = setting.get("result_min", 0)
        r_max = setting.get("result_max", 1000)
        mid_min = setting.get("mid_result_min", 0)  # 中间结果限制

        # 尝试生成符合结果范围的题目
        for _ in range(100):  # 最多重试100次
            a = random.randint(t1_min, t1_max)
            problem_str = str(a)
            current_value = a

            valid = True
            step1_str = ""  # 初始化以修复 linter 错误
            for i in range(steps):
                # 根据步骤选择对应的符号池
                if i == 0:
                    op = random.choice(ops_pool1)
                    b = random.randint(t2_min, t2_max)
                    # 记录第一步的运算过程和结果，用于答案显示
                    step1_val = current_value
                    step1_res = 0  # 初始化
                    if op == "+":
                        step1_res = current_value + b
                    elif op == "-":
                        step1_res = current_value - b
                    elif op == "*":
                        step1_res = current_value * b
                    elif op == "/":
                        step1_res = current_value // b
                    display_op_step1 = op.replace("*", "×").replace("/", "÷")
                    step1_str = f"({step1_val}{display_op_step1}{b}={step1_res})"
                else:
                    op = random.choice(ops_pool2)
                    b = random.randint(t3_min, t3_max)

                if op == "+":
                    current_value += b
                elif op == "-":
                    if current_value < b:
                        valid = False
                        break
                    current_value -= b
                elif op == "*":
                    current_value *= b
                elif op == "/":
                    if b == 0 or current_value % b != 0:
                        valid = False
                        break
                    current_value //= b

                # 检查中间步骤结果是否符合范围（主要防止负数）
                if i < steps - 1:
                    if current_value < mid_min:
                        valid = False
                        break

                # 转换符号显示
                display_op = op.replace("*", "×").replace("/", "÷")
                problem_str += f" {display_op} {b}"

            if valid and r_min <= current_value <= r_max:
                result_text = f"{problem_str} ="
                # 对于2步及以上运算，答案中包含第一步结果
                if steps >= 2:
                    ans_text = f"{problem_str} = {current_value} {step1_str}"
                else:
                    ans_text = f"{problem_str} = {current_value}"
                return result_text, ans_text

        return "1 + 1 =", "1 + 1 = 2"  # 保底方案

    def create_docx(self):
        """生成 Word 文档"""
        quiz_cfg = self.config["quiz"]
        count = quiz_cfg.get("count", 100)
        pages = quiz_cfg.get("pages", 1)
        columns = quiz_cfg.get("columns", 4)
        title = quiz_cfg.get("title", "小学生口算题")
        output_file = quiz_cfg.get("output_file", "小学口算题_v2.docx")
        output_file_answer = quiz_cfg.get("output_file_answer", "小学口算题_答案.docx")

        self.logger.info(
            f"开始生成 {pages} 页，每页 {count} 道题目，并导出到 {output_file} 及 {output_file_answer}"
        )

        doc = Document()
        doc_answer = Document()
        all_unique_problems = set()

        # 读取字体配置
        font_name = quiz_cfg.get("font_name", "黑体")
        font_size = quiz_cfg.get("font_size", 22)
        info_font_size = quiz_cfg.get("info_font_size", 16)
        margin_cm = quiz_cfg.get("margin_cm", 1.0)
        orientation = quiz_cfg.get("orientation", "landscape")

        for page in range(pages):
            if page > 0:
                doc.add_page_break()
                doc_answer.add_page_break()

            # 设置页面布局和标题
            for d in [doc, doc_answer]:
                section = d.sections[-1]
                if orientation == "landscape":
                    new_width, new_height = section.page_height, section.page_width
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width = new_width
                    section.page_height = new_height

                section.top_margin = Cm(margin_cm)
                section.bottom_margin = Cm(margin_cm)
                section.left_margin = Cm(margin_cm)
                section.right_margin = Cm(margin_cm)

                # 添加标题
                current_title = title if d == doc else f"{title} (答案)"
                heading = d.add_heading(current_title, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 题目卷专有信息行
            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_run = info_para.add_run(
                "姓名：__________ 日期：____月____日 时间：________ 对题：____道"
            )
            info_run.font.size = Pt(info_font_size)
            info_run.font.name = font_name
            rPr = info_run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), font_name)

            rows = (count + columns - 1) // columns
            table = doc.add_table(rows=rows, cols=columns)
            table.autofit = True
            table_answer = doc_answer.add_table(rows=rows, cols=columns)
            table_answer.autofit = True

            current_page_problems = []
            current_page_answers = []
            max_retries_per_prob = 1000

            for _ in range(count):
                found_new = False
                for _retry in range(max_retries_per_prob):
                    prob, ans = self.generate_problem()
                    if prob not in all_unique_problems:
                        all_unique_problems.add(prob)
                        current_page_problems.append(prob)
                        current_page_answers.append(ans)
                        found_new = True
                        break

                if not found_new:
                    for _retry in range(max_retries_per_prob):
                        prob, ans = self.generate_problem()
                        if prob not in set(current_page_problems):
                            current_page_problems.append(prob)
                            current_page_answers.append(ans)
                            all_unique_problems.add(prob)
                            found_new = True
                            self.logger.warning(
                                f"警告：第 {page + 1} 页尝试生成全局唯一题目失败，已改用页内唯一模式。"
                            )
                            break

                if not found_new:
                    self.logger.error(
                        f"严重警告：第 {page + 1} 页范围极窄，已无法维持页内题目唯一性。"
                    )
                    prob, ans = self.generate_problem()
                    current_page_problems.append(prob)
                    current_page_answers.append(ans)

            for i in range(len(current_page_problems)):
                row = i // columns
                col = i % columns

                # 填充题目
                cell = table.cell(row, col)
                run = cell.paragraphs[0].add_run(current_page_problems[i])
                run.font.size = Pt(font_size)
                run.font.name = font_name
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn("w:eastAsia"), font_name)

                # 填充答案
                cell_a = table_answer.cell(row, col)
                run_a = cell_a.paragraphs[0].add_run(current_page_answers[i])
                run_a.font.size = Pt(font_size)
                run_a.font.name = font_name
                rPr_a = run_a._element.get_or_add_rPr()
                rFonts_a = rPr_a.get_or_add_rFonts()
                rFonts_a.set(qn("w:eastAsia"), font_name)

        doc.save(output_file)
        doc_answer.save(output_file_answer)
        self.logger.info(f"成功生成文档: {os.path.abspath(output_file)}")
        self.logger.info(f"成功生成答案: {os.path.abspath(output_file_answer)}")


if __name__ == "__main__":
    generator = MathQuizGenerator()
    generator.create_docx()
    convert_to_pdf.convert_docx_to_pdf()
