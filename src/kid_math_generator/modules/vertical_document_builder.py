"""竖式计算题目卷与答案卷的 Word 排版能力。"""

from __future__ import annotations

import math
import random
import tempfile
from collections.abc import Callable, Mapping
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image

from logging_config import get_logger
from kid_math_generator.modules.document_builder import (
    DocumentPair,
    QuizDocumentBuilder,
)
from kid_math_generator.modules.vertical_arithmetic import (
    AdditionWorking,
    SubtractionWorking,
    VerticalArithmeticProblem,
)

VerticalProblemFactory = Callable[[], VerticalArithmeticProblem]
LOGGER = get_logger(__name__)


class VerticalArithmeticDocumentBuilder(QuizDocumentBuilder):
    """把结构化竖式题排版为题目卷和带过程的答案卷。"""

    def __init__(
        self,
        config: Mapping[str, Any],
        *,
        asset_dir: str | Path,
        rng: random.Random | None = None,
    ) -> None:
        super().__init__(config, asset_dir=asset_dir, rng=rng)

    def build(
        self,
        problem_factory: VerticalProblemFactory,
        *,
        output_dir: str | Path,
    ) -> DocumentPair:
        """生成竖式题目卷和答案卷，并返回两个 DOCX 路径。"""
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        question_name = str(self.config.get("output_file", "竖式计算练习.docx"))
        answer_name = str(
            self.config.get("output_file_answer", "竖式计算练习_答案.docx")
        )
        hard_label = bool(self.config.get("hard_label", False))
        if hard_label:
            question_name = self._append_suffix(question_name, "_难题")
            answer_name = self._append_suffix(answer_name, "_难题")

        paths = DocumentPair(
            question=output_path / question_name,
            answer=output_path / answer_name,
        )
        LOGGER.info("生成竖式题目卷 %s 和答案卷 %s", paths.question, paths.answer)

        question_doc = Document()
        answer_doc = Document()
        global_problems: set[tuple[str, int, int]] = set()
        label_image = self._find_first_image() if hard_label else None
        prepared_label = None
        label_temp_dir: tempfile.TemporaryDirectory[str] | None = None

        try:
            if label_image is not None:
                label_temp_dir = tempfile.TemporaryDirectory(prefix="kid_math_label_")
                prepared_label = self._make_background_transparent(
                    Image.open(label_image),
                    int(self.config.get("hard_label_bg_tolerance", 45)),
                )
            elif hard_label:
                LOGGER.warning("已启用 hard_label，但 %s 中未找到图片", self.asset_dir)

            pages = max(1, int(self.config.get("pages", 1)))
            count = max(1, int(self.config.get("count", 12)))
            columns = max(1, int(self.config.get("columns", 3)))

            for page_number in range(pages):
                if page_number:
                    question_doc.add_page_break()
                    answer_doc.add_page_break()

                title = str(self.config.get("title", "竖式计算练习"))
                question_heading = self._add_page_header(
                    question_doc,
                    title,
                    include_info=True,
                )
                self._add_page_header(
                    answer_doc,
                    f"{title}（答案）",
                    include_info=False,
                )

                if prepared_label is not None and label_temp_dir is not None:
                    self._add_random_label(
                        question_heading,
                        prepared_label,
                        Path(label_temp_dir.name),
                        page_number,
                    )

                problems = self._generate_page(
                    problem_factory,
                    count,
                    global_problems,
                    page_number,
                )
                self._add_vertical_problem_tables(
                    question_doc,
                    answer_doc,
                    problems,
                    columns,
                )

            question_doc.save(paths.question)
            answer_doc.save(paths.answer)
        finally:
            if label_temp_dir is not None:
                label_temp_dir.cleanup()

        return paths

    def _generate_page(
        self,
        problem_factory: VerticalProblemFactory,
        count: int,
        global_problems: set[tuple[str, int, int]],
        page_number: int,
    ) -> list[VerticalArithmeticProblem]:
        problems: list[VerticalArithmeticProblem] = []
        page_keys: set[tuple[str, int, int]] = set()

        for _ in range(count):
            selected: VerticalArithmeticProblem | None = None
            for require_global_unique in (True, False):
                for _attempt in range(2000):
                    problem = problem_factory()
                    key = problem.unique_key
                    if key in page_keys:
                        continue
                    if require_global_unique and key in global_problems:
                        continue
                    selected = problem
                    break
                if selected is not None:
                    if not require_global_unique:
                        LOGGER.warning("第 %s 页已退回页内唯一模式", page_number + 1)
                    break

            if selected is None:
                selected = problem_factory()
                LOGGER.warning("第 %s 页题目范围不足，允许重复题目", page_number + 1)

            problems.append(selected)
            page_keys.add(selected.unique_key)
            global_problems.add(selected.unique_key)

        return problems

    def _add_vertical_problem_tables(
        self,
        question_doc: Document,
        answer_doc: Document,
        problems: list[VerticalArithmeticProblem],
        columns: int,
    ) -> None:
        rows = math.ceil(len(problems) / columns)
        question_table = question_doc.add_table(rows=rows, cols=columns)
        answer_table = answer_doc.add_table(rows=rows, cols=columns)
        self._setup_table(question_table, question_doc.sections[-1], columns)
        self._setup_table(
            answer_table,
            answer_doc.sections[-1],
            columns,
            is_answer=True,
        )

        show_working = bool(self.config.get("show_working_in_answer", True))
        section = question_doc.sections[-1]
        available_width = (
            section.page_width - section.left_margin - section.right_margin
        )
        container_width = available_width // columns
        for index, problem in enumerate(problems):
            row, column = divmod(index, columns)
            self._render_problem(
                question_table.cell(row, column),
                problem,
                number=index + 1,
                show_answer=False,
                show_working=False,
                container_width=container_width,
            )
            self._render_problem(
                answer_table.cell(row, column),
                problem,
                number=index + 1,
                show_answer=True,
                show_working=show_working,
                container_width=container_width,
            )

    def _render_problem(
        self,
        outer_cell,
        problem: VerticalArithmeticProblem,
        *,
        number: int,
        show_answer: bool,
        show_working: bool,
        container_width: int,
    ) -> None:
        number_paragraph = outer_cell.paragraphs[0]
        number_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        number_run = number_paragraph.add_run(f"{number}.")
        self._set_run_font(
            number_run,
            str(self.config.get("font_name", "黑体")),
            int(self.config.get("number_font_size", 11)),
        )

        width = max(
            len(str(problem.left)),
            len(str(problem.right)),
            len(str(problem.result)),
        )
        digit_start = 2
        table = outer_cell.add_table(rows=4, cols=width + 3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._remove_table_borders(table)

        operator_width = Cm(0.55)
        digit_width = Cm(float(self.config.get("digit_cell_width_cm", 0.65)))
        compact_width = int(operator_width) + int(digit_width) * width
        inner_width = max(compact_width, container_width - int(Cm(0.25)))
        spacer_width = max(1, (inner_width - compact_width) // 2)
        trailing_width = inner_width - compact_width - spacer_width
        self._set_fixed_table_geometry(
            table,
            [
                spacer_width,
                int(operator_width),
                *([int(digit_width)] * width),
                trailing_width,
            ],
        )
        for row in table.rows:
            for column_index, cell in enumerate(row.cells):
                if column_index == 0:
                    cell.width = spacer_width
                elif column_index == 1:
                    cell.width = operator_width
                elif column_index == width + 2:
                    cell.width = trailing_width
                else:
                    cell.width = digit_width
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                self._set_cell_margins(
                    cell,
                    top=0,
                    start=15,
                    bottom=0,
                    end=15,
                )
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1

        self._write_working_row(
            table,
            problem,
            width,
            digit_start,
            show_working,
        )
        self._write_number_row(table, 1, problem.left, width, digit_start)
        self._write_number_row(
            table,
            2,
            problem.right,
            width,
            digit_start,
            operator=problem.symbol,
        )
        for cell in table.rows[2].cells[1 : digit_start + width]:
            self._set_bottom_border(cell)
            self._set_paragraph_bottom_border(cell.paragraphs[0])
        if show_answer:
            self._write_number_row(table, 3, problem.result, width, digit_start)
        else:
            self._ensure_blank_row_height(table, 3)

    def _write_working_row(
        self,
        table,
        problem: VerticalArithmeticProblem,
        width: int,
        digit_start: int,
        show_working: bool,
    ) -> None:
        markers: list[str] = [""] * width
        if show_working and isinstance(problem.working, AdditionWorking):
            markers = [str(value) if value is not None else "" for value in problem.working.carries]
        elif show_working and isinstance(problem.working, SubtractionWorking):
            markers = ["借" if value else "" for value in problem.working.borrows]

        for index, marker in enumerate(markers, start=digit_start):
            self._write_cell_text(table.cell(0, index), marker, working=True)
        self._ensure_blank_row_height(table, 0)

    def _write_number_row(
        self,
        table,
        row_index: int,
        value: int,
        width: int,
        digit_start: int,
        *,
        operator: str = "",
    ) -> None:
        self._write_cell_text(table.cell(row_index, digit_start - 1), operator)
        padded = str(value).rjust(width)
        for index, character in enumerate(padded, start=digit_start):
            self._write_cell_text(table.cell(row_index, index), character.strip())

    def _write_cell_text(self, cell, text: str, *, working: bool = False) -> None:
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        self._set_run_font(
            run,
            str(self.config.get("font_name", "黑体")),
            int(
                self.config.get(
                    "working_font_size" if working else "font_size",
                    9 if working else 20,
                )
            ),
        )

    def _ensure_blank_row_height(self, table, row_index: int) -> None:
        height = Cm(float(self.config.get("digit_row_height_cm", 0.65)))
        table.rows[row_index].height = height

    @staticmethod
    def _remove_table_borders(table) -> None:
        properties = table._tbl.tblPr
        borders = properties.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            properties.append(borders)
        for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = borders.find(qn(f"w:{name}"))
            if border is None:
                border = OxmlElement(f"w:{name}")
                borders.append(border)
            border.set(qn("w:val"), "nil")

    @staticmethod
    def _set_bottom_border(cell) -> None:
        properties = cell._tc.get_or_add_tcPr()
        borders = properties.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            properties.append(borders)
        bottom = borders.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            borders.append(bottom)
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:color"), "000000")

    @staticmethod
    def _set_paragraph_bottom_border(paragraph) -> None:
        properties = paragraph._p.get_or_add_pPr()
        borders = properties.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            properties.append(borders)
        bottom = borders.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            borders.append(bottom)
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
