"""口算题与答案卷的 Word 文档生成能力。"""

from __future__ import annotations

import math
import random
import tempfile
from collections.abc import Callable, Mapping
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image

from logging_config import get_logger

ProblemFactory = Callable[[], tuple[str, str]]
LOGGER = get_logger(__name__)


@dataclass(frozen=True, slots=True)
class DocumentPair:
    question: Path
    answer: Path


class QuizDocumentBuilder:
    """把任意题目生成器的结构化结果排版成题目卷和答案卷。"""

    A4_WIDTH_CM = 21.0
    A4_HEIGHT_CM = 29.7
    PT_TO_TWIPS = 20
    CM_TO_PT = 28.3465
    EMU_PER_TWIP = 635

    def __init__(
        self,
        config: Mapping[str, Any],
        *,
        asset_dir: str | Path,
        rng: random.Random | None = None,
    ) -> None:
        self.config = dict(config)
        self.asset_dir = Path(asset_dir)
        self.rng = rng or random.Random()

    def build(
        self,
        problem_factory: ProblemFactory,
        *,
        output_dir: str | Path,
    ) -> DocumentPair:
        """生成题目卷和答案卷，并返回两个 DOCX 路径。"""
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        question_name = str(self.config.get("output_file", "口算题.docx"))
        answer_name = str(self.config.get("output_file_answer", "口算题_答案.docx"))
        hard_label = bool(self.config.get("hard_label", False))
        if hard_label:
            question_name = self._append_suffix(question_name, "_难题")
            answer_name = self._append_suffix(answer_name, "_难题")

        paths = DocumentPair(
            question=output_path / question_name,
            answer=output_path / answer_name,
        )
        LOGGER.info("生成题目卷 %s 和答案卷 %s", paths.question, paths.answer)

        question_doc = Document()
        answer_doc = Document()
        global_problems: set[str] = set()
        label_enabled = bool(self.config.get("label_enabled", hard_label))
        label_images = self._find_images() if label_enabled else ()
        prepared_labels: list[Image.Image] = []
        label_temp_dir: tempfile.TemporaryDirectory[str] | None = None

        try:
            if label_images:
                label_temp_dir = tempfile.TemporaryDirectory(prefix="kid_math_label_")
                for label_image in label_images:
                    with Image.open(label_image) as source_image:
                        prepared_labels.append(
                            self._make_background_transparent(
                                source_image,
                                int(self.config.get("hard_label_bg_tolerance", 45)),
                            )
                        )
            elif label_enabled:
                LOGGER.warning("已启用盖章，但 %s 中未找到图片", self.asset_dir)

            pages = max(1, int(self.config.get("pages", 1)))
            count = max(1, int(self.config.get("count", 20)))
            columns = max(1, int(self.config.get("columns", 2)))

            for page_number in range(pages):
                if page_number:
                    question_doc.add_page_break()
                    answer_doc.add_page_break()

                question_heading = self._add_page_header(
                    question_doc,
                    str(self.config.get("title", "小学生口算题")),
                    include_info=True,
                )
                self._add_page_header(
                    answer_doc,
                    f"{self.config.get('title', '小学生口算题')}（答案）",
                    include_info=False,
                )

                if prepared_labels and label_temp_dir is not None:
                    self._add_random_label(
                        question_heading,
                        self.rng.choice(prepared_labels),
                        Path(label_temp_dir.name),
                        page_number,
                    )

                problems, answers = self._generate_page(
                    problem_factory,
                    count,
                    global_problems,
                    page_number,
                )
                self._add_problem_tables(
                    question_doc,
                    answer_doc,
                    problems,
                    answers,
                    columns,
                )

            question_doc.save(paths.question)
            answer_doc.save(paths.answer)
        finally:
            if label_temp_dir is not None:
                label_temp_dir.cleanup()

        return paths

    def _add_page_header(
        self,
        document: Document,
        title: str,
        *,
        include_info: bool,
    ):
        section = document.sections[-1]
        self._apply_page_layout(section)
        heading = document.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if include_info:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(
                "姓名：__________ 日期：____月____日 时间：________ 对题：____道"
            )
            self._set_run_font(
                run,
                str(self.config.get("font_name", "黑体")),
                int(self.config.get("info_font_size", 16)),
            )
        return heading

    def _generate_page(
        self,
        problem_factory: ProblemFactory,
        count: int,
        global_problems: set[str],
        page_number: int,
    ) -> tuple[list[str], list[str]]:
        problems: list[str] = []
        answers: list[str] = []
        page_problems: set[str] = set()

        for _ in range(count):
            selected: tuple[str, str] | None = None
            for require_global_unique in (True, False):
                for _attempt in range(2000):
                    problem, answer = problem_factory()
                    if problem in page_problems:
                        continue
                    if require_global_unique and problem in global_problems:
                        continue
                    selected = problem, answer
                    break
                if selected is not None:
                    if not require_global_unique:
                        LOGGER.warning("第 %s 页已退回页内唯一模式", page_number + 1)
                    break

            if selected is None:
                problem, answer = problem_factory()
                LOGGER.warning("第 %s 页题目范围不足，允许重复题目", page_number + 1)
            else:
                problem, answer = selected

            problems.append(problem)
            answers.append(answer)
            page_problems.add(problem)
            global_problems.add(problem)

        return problems, answers

    def _add_problem_tables(
        self,
        question_doc: Document,
        answer_doc: Document,
        problems: list[str],
        answers: list[str],
        columns: int,
    ) -> None:
        rows = math.ceil(len(problems) / columns)
        question_table = question_doc.add_table(rows=rows, cols=columns)
        answer_table = answer_doc.add_table(rows=rows, cols=columns)
        self._setup_table(question_table, question_doc.sections[-1], columns)
        self._setup_table(
            answer_table,
            answer_doc.sections[-1],
            columns,
            is_answer=True,
        )

        answer_font_size = self._calculate_answer_font_size(
            answer_doc.sections[-1],
            rows,
            columns,
            answers,
        )
        font_name = str(self.config.get("font_name", "黑体"))
        problem_font_size = int(self.config.get("font_size", 22))

        for index, (problem, answer) in enumerate(zip(problems, answers)):
            row, column = divmod(index, columns)
            problem_run = question_table.cell(row, column).paragraphs[0].add_run(problem)
            answer_run = answer_table.cell(row, column).paragraphs[0].add_run(answer)
            self._set_run_font(problem_run, font_name, problem_font_size)
            self._set_run_font(answer_run, font_name, answer_font_size)

    def _setup_table(
        self,
        table,
        section,
        columns: int,
        *,
        is_answer: bool = False,
    ) -> None:
        width = section.page_width - section.left_margin - section.right_margin
        column_widths = [width // columns] * columns
        column_widths[-1] += width - sum(column_widths)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_fixed_table_geometry(table, column_widths)

        for column_index, column in enumerate(table.columns):
            for cell in column.cells:
                cell.width = column_widths[column_index]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                self._set_cell_margins(
                    cell,
                    top=20 if is_answer else 40,
                    start=40 if is_answer else 60,
                    bottom=20 if is_answer else 40,
                    end=40 if is_answer else 60,
                )
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1

    @classmethod
    def _set_fixed_table_geometry(cls, table, column_widths: list[int]) -> None:
        """以 Word 使用的 DXA 单位设置表格、网格列和单元格宽度。"""
        widths_twips = [
            max(1, round(width / cls.EMU_PER_TWIP))
            for width in column_widths
        ]
        total_twips = sum(widths_twips)
        properties = table._tbl.tblPr

        table_width = properties.first_child_found_in("w:tblW")
        if table_width is None:
            table_width = OxmlElement("w:tblW")
            properties.append(table_width)
        table_width.set(qn("w:w"), str(total_twips))
        table_width.set(qn("w:type"), "dxa")

        layout = properties.first_child_found_in("w:tblLayout")
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            properties.append(layout)
        layout.set(qn("w:type"), "fixed")

        grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
        for grid_column, width_twips in zip(grid_columns, widths_twips):
            grid_column.set(qn("w:w"), str(width_twips))

        for row in table.rows:
            for cell, width_twips in zip(row.cells, widths_twips):
                cell_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
                if cell_width is None:
                    cell_width = OxmlElement("w:tcW")
                    cell._tc.get_or_add_tcPr().append(cell_width)
                cell_width.set(qn("w:w"), str(width_twips))
                cell_width.set(qn("w:type"), "dxa")

    def _calculate_answer_font_size(
        self,
        section,
        rows: int,
        columns: int,
        answers: list[str],
    ) -> int:
        preferred = int(
            self.config.get(
                "answer_font_size",
                max(12, int(self.config.get("font_size", 22)) - 4),
            )
        )
        if not answers:
            return preferred

        margin_pt = float(self.config.get("margin_cm", 1.0)) * self.CM_TO_PT
        width_pt = section.page_width / self.PT_TO_TWIPS - 2 * margin_pt
        height_pt = section.page_height / self.PT_TO_TWIPS - 2 * margin_pt
        column_width_pt = max(width_pt, 100) / columns
        row_height_pt = max(height_pt - 70, 60) / rows
        max_length = max(len(answer) for answer in answers)
        width_size = column_width_pt / max(max_length * 0.9, 1)
        height_size = max((row_height_pt - 6) / 1.35, 1)
        return max(8, int(min(preferred, width_size, height_size)))

    def _apply_page_layout(self, section) -> None:
        orientation = str(self.config.get("orientation", "landscape")).lower()
        if orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Cm(self.A4_HEIGHT_CM)
            section.page_height = Cm(self.A4_WIDTH_CM)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Cm(self.A4_WIDTH_CM)
            section.page_height = Cm(self.A4_HEIGHT_CM)

        margin = Cm(float(self.config.get("margin_cm", 1.0)))
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

    def _find_first_image(self) -> Path | None:
        images = self._find_images()
        return images[0] if images else None

    def _find_images(self) -> tuple[Path, ...]:
        """返回素材目录下可用于盖章的全部图片。"""
        if not self.asset_dir.is_dir():
            return ()
        supported = {".png", ".jpg", ".jpeg", ".bmp", ".gif"}
        return tuple(
            path
            for path in sorted(self.asset_dir.iterdir())
            if path.is_file() and path.suffix.lower() in supported
        )

    def _add_random_label(
        self,
        heading,
        image: Image.Image,
        temp_dir: Path,
        page_number: int,
    ) -> None:
        rotation_min = float(self.config.get("hard_label_rotation_min", -15))
        rotation_max = float(self.config.get("hard_label_rotation_max", 15))
        angle = self.rng.uniform(rotation_min, rotation_max)
        resampling_group = getattr(Image, "Resampling", Image)
        maximum_width = max(
            1,
            int(self.config.get("hard_label_max_width_px", 600)),
        )
        prepared = image.copy()
        if prepared.width > maximum_width:
            height = max(1, round(prepared.height * maximum_width / prepared.width))
            prepared = prepared.resize(
                (maximum_width, height),
                resample=getattr(resampling_group, "LANCZOS"),
            )
        rotated = prepared.rotate(
            angle,
            expand=True,
            resample=getattr(resampling_group, "BICUBIC"),
            fillcolor=(255, 255, 255, 0),
        )
        image_path = temp_dir / f"hard_label_page_{page_number + 1}.png"
        rotated.save(image_path)

        x = max(
            0.0,
            float(self.config.get("hard_label_offset_x_cm", 0))
            + self.rng.uniform(
                -float(self.config.get("hard_label_jitter_x_cm", 0.45)),
                float(self.config.get("hard_label_jitter_x_cm", 0.45)),
            ),
        )
        y = max(
            0.0,
            float(self.config.get("hard_label_offset_y_cm", 0))
            + self.rng.uniform(
                -float(self.config.get("hard_label_jitter_y_cm", 0.3)),
                float(self.config.get("hard_label_jitter_y_cm", 0.3)),
            ),
        )
        run = heading.add_run()
        run.add_picture(
            str(image_path),
            width=Cm(float(self.config.get("hard_label_width_cm", 2.2))),
        )
        self._make_picture_float(run, x, y)

    @staticmethod
    def _make_background_transparent(
        image: Image.Image,
        tolerance: int,
    ) -> Image.Image:
        rgba = image.convert("RGBA")
        neutral_tolerance = min(max(tolerance, 0), 24)
        pixels = []
        pixel_data = (
            rgba.get_flattened_data()
            if hasattr(rgba, "get_flattened_data")
            else rgba.getdata()
        )
        for red, green, blue, alpha in pixel_data:
            channels = (red, green, blue)
            neutral_noise = max(channels) - min(channels) <= neutral_tolerance
            pixels.append(
                (red, green, blue, 0)
                if neutral_noise
                else (red, green, blue, alpha)
            )
        rgba.putdata(pixels)
        return rgba

    @staticmethod
    def _make_picture_float(run, x_cm: float, y_cm: float) -> None:
        anchor = run._r.xpath(".//wp:inline")[0]
        anchor.tag = qn("wp:anchor")
        for name, value in {
            "distT": "0",
            "distB": "0",
            "distL": "0",
            "distR": "0",
            "simplePos": "0",
            "relativeHeight": "251658240",
            "behindDoc": "0",
            "locked": "0",
            "layoutInCell": "1",
            "allowOverlap": "1",
        }.items():
            anchor.set(name, value)

        simple_pos = OxmlElement("wp:simplePos")
        simple_pos.set("x", "0")
        simple_pos.set("y", "0")
        position_h = OxmlElement("wp:positionH")
        position_h.set("relativeFrom", "page")
        position_h_offset = OxmlElement("wp:posOffset")
        position_h_offset.text = str(int(Cm(x_cm)))
        position_h.append(position_h_offset)
        position_v = OxmlElement("wp:positionV")
        position_v.set("relativeFrom", "page")
        position_v_offset = OxmlElement("wp:posOffset")
        position_v_offset.text = str(int(Cm(y_cm)))
        position_v.append(position_v_offset)

        anchor.insert(0, simple_pos)
        anchor.insert(1, position_h)
        anchor.insert(2, position_v)
        extent = anchor.find(qn("wp:extent"))
        if extent is not None:
            effect_extent = OxmlElement("wp:effectExtent")
            for side in ("l", "t", "r", "b"):
                effect_extent.set(side, "0")
            extent_index = list(anchor).index(extent)
            anchor.insert(extent_index + 1, effect_extent)
            anchor.insert(extent_index + 2, OxmlElement("wp:wrapNone"))

    @staticmethod
    def _set_cell_margins(
        cell,
        *,
        top: int,
        start: int,
        bottom: int,
        end: int,
    ) -> None:
        properties = cell._tc.get_or_add_tcPr()
        margins = properties.first_child_found_in("w:tcMar")
        if margins is None:
            margins = OxmlElement("w:tcMar")
            properties.append(margins)
        for name, value in {
            "top": top,
            "start": start,
            "bottom": bottom,
            "end": end,
        }.items():
            node = margins.find(qn(f"w:{name}"))
            if node is None:
                node = OxmlElement(f"w:{name}")
                margins.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    @staticmethod
    def _set_run_font(run, name: str, size: int) -> None:
        run.font.name = name
        run.font.size = Pt(size)
        fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        fonts.set(qn("w:eastAsia"), name)

    @staticmethod
    def _append_suffix(filename: str, suffix: str) -> str:
        path = Path(filename)
        if path.stem.endswith(suffix):
            return str(path)
        return str(path.with_name(f"{path.stem}{suffix}{path.suffix}"))
