"""配置加载和文档构建测试。"""

from __future__ import annotations

import io
import random
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.shared import Cm
from PIL import Image

from kid_math_generator.config_loader import load_config
from kid_math_generator.modules.document_builder import QuizDocumentBuilder
from kid_math_generator.modules.multiplication import (
    MultiplicationTableProblemGenerator,
)


class ConfigTests(unittest.TestCase):
    def test_project_config_contains_all_flows(self) -> None:
        config = load_config(PROJECT_ROOT / "config.yaml")
        self.assertIn("addition_subtraction", config["flows"])
        self.assertIn("vertical_arithmetic", config["flows"])
        multiplication = config["flows"]["multiplication"]
        self.assertEqual(multiplication["factor_min"], 1)
        self.assertEqual(multiplication["factor_max"], 9)
        self.assertTrue(multiplication["label_enabled"])
        for flow in config["flows"].values():
            self.assertEqual(flow.get("orientation", "landscape"), "landscape")


class DocumentBuilderTests(unittest.TestCase):
    def test_stamps_multiplication_question_pages_at_top_left(self) -> None:
        problems = iter(
            [
                ("2 × 3 =", "2 × 3 = 6"),
                ("4 × 5 =", "4 × 5 = 20"),
            ]
        )

        with tempfile.TemporaryDirectory() as temp_dir:
            directory = Path(temp_dir)
            asset_dir = directory / "assets"
            asset_dir.mkdir()
            Image.new("RGB", (800, 400), (220, 80, 100)).save(asset_dir / "a.png")
            Image.new("RGB", (800, 400), (80, 120, 220)).save(asset_dir / "b.png")
            builder = QuizDocumentBuilder(
                {
                    "pages": 2,
                    "count": 1,
                    "columns": 1,
                    "title": "乘法盖章测试",
                    "output_file": "multiplication.docx",
                    "output_file_answer": "multiplication-answer.docx",
                    "label_enabled": True,
                    "hard_label": False,
                    "hard_label_width_cm": 4.2,
                    "hard_label_max_width_px": 600,
                    "hard_label_offset_x_cm": 0.7,
                    "hard_label_offset_y_cm": 0.3,
                    "hard_label_rotation_min": 0,
                    "hard_label_rotation_max": 0,
                    "hard_label_jitter_x_cm": 0,
                    "hard_label_jitter_y_cm": 0,
                },
                asset_dir=asset_dir,
                rng=random.Random(2),
            )

            pair = builder.build(lambda: next(problems), output_dir=directory)

            self.assertEqual(pair.question.name, "multiplication.docx")
            self.assertEqual(pair.answer.name, "multiplication-answer.docx")
            question_doc = Document(pair.question)
            answer_doc = Document(pair.answer)
            anchors = question_doc.element.body.xpath(".//wp:anchor")
            self.assertEqual(len(anchors), 2)
            self.assertEqual(len(answer_doc.element.body.xpath(".//wp:anchor")), 0)
            for anchor in anchors:
                horizontal = anchor.find(qn("wp:positionH"))
                vertical = anchor.find(qn("wp:positionV"))
                self.assertEqual(
                    int(horizontal.find(qn("wp:posOffset")).text),
                    int(Cm(0.7)),
                )
                self.assertEqual(
                    int(vertical.find(qn("wp:posOffset")).text),
                    int(Cm(0.3)),
                )

            with zipfile.ZipFile(pair.question) as package:
                media = [
                    name for name in package.namelist() if name.startswith("word/media/")
                ]
                media_sizes = {
                    Image.open(io.BytesIO(package.read(name))).size for name in media
                }
            with zipfile.ZipFile(pair.answer) as package:
                answer_media = [
                    name for name in package.namelist() if name.startswith("word/media/")
                ]
            self.assertEqual(media_sizes, {(600, 300)})
            self.assertEqual(answer_media, [])

    def test_builds_question_and_answer_documents(self) -> None:
        config = {
            "pages": 1,
            "count": 6,
            "columns": 2,
            "title": "乘法测试",
            "output_file": "questions.docx",
            "output_file_answer": "answers.docx",
            "font_name": "黑体",
            "font_size": 18,
            "hard_label": False,
        }
        rng = random.Random(3)
        problem_generator = MultiplicationTableProblemGenerator(rng=rng)
        builder = QuizDocumentBuilder(
            config,
            asset_dir=PROJECT_ROOT / "src",
            rng=rng,
        )

        with tempfile.TemporaryDirectory() as temp_dir:
            result = builder.build(
                problem_generator.generate,
                output_dir=temp_dir,
            )
            self.assertTrue(result.question.is_file())
            self.assertTrue(result.answer.is_file())

            question_doc = Document(result.question)
            answer_doc = Document(result.answer)
            for document in (question_doc, answer_doc):
                section = document.sections[0]
                self.assertEqual(section.orientation, WD_ORIENT.LANDSCAPE)
                self.assertAlmostEqual(section.page_width.cm, 29.7, places=1)
                self.assertAlmostEqual(section.page_height.cm, 21.0, places=1)
            question_text = "\n".join(
                cell.text
                for table in question_doc.tables
                for row in table.rows
                for cell in row.cells
            )
            answer_text = "\n".join(
                cell.text
                for table in answer_doc.tables
                for row in table.rows
                for cell in row.cells
            )
            self.assertEqual(question_text.count("="), 6)
            self.assertEqual(answer_text.count("="), 6)
            self.assertIn("×", question_text)


if __name__ == "__main__":
    unittest.main()
