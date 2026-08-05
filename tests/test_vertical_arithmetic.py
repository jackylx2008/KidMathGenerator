"""竖式计算生成规则与 Word 排版测试。"""

from __future__ import annotations

import io
import random
import logging
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.shared import Cm
from PIL import Image

from kid_math_generator.context import AppContext
from kid_math_generator.flows.vertical_arithmetic_flow import run
from kid_math_generator.modules.vertical_arithmetic import (
    AdditionWorking,
    SubtractionWorking,
    VerticalArithmeticProblem,
    VerticalArithmeticProblemGenerator,
)
from kid_math_generator.modules.vertical_document_builder import (
    VerticalArithmeticDocumentBuilder,
)


class VerticalArithmeticGeneratorTests(unittest.TestCase):
    def test_addition_can_require_exactly_one_carry(self) -> None:
        generator = VerticalArithmeticProblemGenerator(
            [
                {
                    "operation": "addition",
                    "left_min": 10,
                    "left_max": 99,
                    "right_min": 10,
                    "right_max": 99,
                    "carry": "required",
                    "carry_count_min": 1,
                    "carry_count_max": 1,
                }
            ],
            rng=random.Random(7),
        )

        for _ in range(100):
            problem = generator.generate()
            self.assertEqual(problem.result, problem.left + problem.right)
            self.assertIsInstance(problem.working, AdditionWorking)
            carry_count = sum(
                carry is not None for carry in problem.working.carries
            )
            self.assertEqual(carry_count, 1)

    def test_addition_can_forbid_carry(self) -> None:
        generator = VerticalArithmeticProblemGenerator(
            [
                {
                    "operation": "+",
                    "left_min": 10,
                    "left_max": 99,
                    "right_min": 10,
                    "right_max": 99,
                    "carry": "none",
                }
            ],
            rng=random.Random(11),
        )

        for _ in range(50):
            problem = generator.generate()
            self.assertTrue(
                all(carry is None for carry in problem.working.carries)
            )

    def test_subtraction_requires_borrow_and_nonnegative_result(self) -> None:
        generator = VerticalArithmeticProblemGenerator(
            [
                {
                    "operation": "subtraction",
                    "left_min": 10,
                    "left_max": 99,
                    "right_min": 10,
                    "right_max": 99,
                    "borrow": "required",
                    "allow_negative": False,
                }
            ],
            rng=random.Random(13),
        )

        for _ in range(100):
            problem = generator.generate()
            self.assertEqual(problem.result, problem.left - problem.right)
            self.assertGreaterEqual(problem.result, 0)
            self.assertIsInstance(problem.working, SubtractionWorking)
            self.assertTrue(any(problem.working.borrows))

    def test_future_operation_fails_with_explicit_message(self) -> None:
        with self.assertRaisesRegex(ValueError, "暂未实现.*multiplication"):
            VerticalArithmeticProblemGenerator(
                [{"operation": "multiplication"}],
            )


class VerticalArithmeticDocumentTests(unittest.TestCase):
    def test_randomly_stamps_question_pages_at_top_left(self) -> None:
        problems = iter(
            [
                VerticalArithmeticProblem(
                    operation="addition",
                    left=47,
                    right=28,
                    result=75,
                    working=AdditionWorking((1, None)),
                ),
                VerticalArithmeticProblem(
                    operation="subtraction",
                    left=52,
                    right=38,
                    result=14,
                    working=SubtractionWorking((False, True)),
                ),
            ]
        )

        with tempfile.TemporaryDirectory() as temp_dir:
            directory = Path(temp_dir)
            asset_dir = directory / "assets"
            asset_dir.mkdir()
            Image.new("RGB", (800, 400), (220, 80, 100)).save(asset_dir / "a.png")
            Image.new("RGB", (800, 400), (80, 120, 220)).save(asset_dir / "b.png")
            builder = VerticalArithmeticDocumentBuilder(
                {
                    "pages": 2,
                    "count": 1,
                    "columns": 1,
                    "title": "盖章测试",
                    "output_file": "stamped.docx",
                    "output_file_answer": "stamped-answer.docx",
                    "label_enabled": True,
                    "hard_label": False,
                    "hard_label_width_cm": 4.2,
                    "hard_label_offset_x_cm": 0.7,
                    "hard_label_offset_y_cm": 0.3,
                    "hard_label_rotation_min": 0,
                    "hard_label_rotation_max": 0,
                    "hard_label_jitter_x_cm": 0,
                    "hard_label_jitter_y_cm": 0,
                },
                asset_dir=asset_dir,
                rng=random.Random(2),
            )

            pair = builder.build(lambda: next(problems), output_dir=directory)

            self.assertEqual(pair.question.name, "stamped.docx")
            self.assertEqual(pair.answer.name, "stamped-answer.docx")
            question_doc = Document(pair.question)
            answer_doc = Document(pair.answer)
            anchors = question_doc.element.body.xpath(".//wp:anchor")
            self.assertEqual(len(anchors), 2)
            self.assertEqual(len(answer_doc.element.body.xpath(".//wp:anchor")), 0)
            for anchor in anchors:
                horizontal = anchor.find(qn("wp:positionH"))
                vertical = anchor.find(qn("wp:positionV"))
                self.assertEqual(
                    int(horizontal.find(qn("wp:posOffset")).text),
                    int(Cm(0.7)),
                )
                self.assertEqual(
                    int(vertical.find(qn("wp:posOffset")).text),
                    int(Cm(0.3)),
                )

            with zipfile.ZipFile(pair.question) as package:
                media = [
                    name for name in package.namelist() if name.startswith("word/media/")
                ]
                media_sizes = [
                    Image.open(io.BytesIO(package.read(name))).size for name in media
                ]
            with zipfile.ZipFile(pair.answer) as package:
                answer_media = [
                    name for name in package.namelist() if name.startswith("word/media/")
                ]
            self.assertEqual(len(media), 2)
            self.assertEqual(media_sizes, [(600, 300), (600, 300)])
            self.assertEqual(answer_media, [])

    def test_stamp_background_cleanup_preserves_colored_artwork(self) -> None:
        image = Image.new("RGBA", (3, 1))
        image.putdata(
            [
                (0, 0, 0, 255),
                (255, 255, 255, 255),
                (220, 80, 100, 255),
            ]
        )

        cleaned = VerticalArithmeticDocumentBuilder._make_background_transparent(
            image,
            45,
        )

        self.assertEqual(cleaned.getpixel((0, 0))[3], 0)
        self.assertEqual(cleaned.getpixel((1, 0))[3], 0)
        self.assertEqual(cleaned.getpixel((2, 0))[3], 255)

    def test_builds_digit_aligned_question_and_answer_documents(self) -> None:
        problems = iter(
            [
                VerticalArithmeticProblem(
                    operation="addition",
                    left=47,
                    right=28,
                    result=75,
                    working=AdditionWorking((1, None)),
                ),
                VerticalArithmeticProblem(
                    operation="subtraction",
                    left=52,
                    right=38,
                    result=14,
                    working=SubtractionWorking((False, True)),
                ),
            ]
        )
        builder = VerticalArithmeticDocumentBuilder(
            {
                "pages": 1,
                "count": 2,
                "columns": 2,
                "title": "竖式测试",
                "output_file": "vertical.docx",
                "output_file_answer": "vertical-answer.docx",
                "font_name": "黑体",
                "font_size": 18,
                "show_working_in_answer": True,
                "hard_label": False,
            },
            asset_dir=PROJECT_ROOT / "src",
            rng=random.Random(3),
        )

        with tempfile.TemporaryDirectory() as temp_dir:
            pair = builder.build(lambda: next(problems), output_dir=temp_dir)
            question_doc = Document(pair.question)
            answer_doc = Document(pair.answer)

            self.assertEqual(len(question_doc.tables), 1)
            self.assertEqual(len(answer_doc.tables), 1)
            for document in (question_doc, answer_doc):
                section = document.sections[0]
                self.assertEqual(section.orientation, WD_ORIENT.LANDSCAPE)
                self.assertAlmostEqual(section.page_width.cm, 29.7, places=1)
                self.assertAlmostEqual(section.page_height.cm, 21.0, places=1)
                table_tail = document.paragraphs[-1]
                self.assertEqual(table_tail.text, "")
                self.assertEqual(table_tail.paragraph_format.line_spacing.pt, 1.0)
                self.assertEqual(table_tail.paragraph_format.space_before.pt, 0.0)
                self.assertEqual(table_tail.paragraph_format.space_after.pt, 0.0)
            info_paragraph = question_doc.paragraphs[1]
            self.assertEqual(info_paragraph.paragraph_format.space_before.pt, 0.0)
            self.assertEqual(info_paragraph.paragraph_format.space_after.pt, 0.0)
            outer_width = question_doc.tables[0]._tbl.tblPr.first_child_found_in(
                "w:tblW"
            )
            self.assertEqual(outer_width.get(qn("w:type")), "dxa")
            self.assertGreater(int(outer_width.get(qn("w:w"))), 15000)
            self.assertLess(int(outer_width.get(qn("w:w"))), 16000)
            question_cells = question_doc.tables[0].rows[0].cells
            answer_cells = answer_doc.tables[0].rows[0].cells
            question_verticals = [cell.tables[0] for cell in question_cells]
            answer_verticals = [cell.tables[0] for cell in answer_cells]
            for cell in (*question_cells, *answer_cells):
                cell_tail = cell.paragraphs[-1]
                self.assertEqual(cell_tail.text, "")
                self.assertEqual(cell_tail.paragraph_format.line_spacing.pt, 1.0)

            self.assertEqual(self._row_text(question_verticals[0], 1), "47")
            self.assertEqual(self._row_text(question_verticals[0], 2), "+28")
            self.assertEqual(self._row_text(question_verticals[0], 3), "")
            self.assertEqual(self._row_text(answer_verticals[0], 0), "1")
            self.assertEqual(self._row_text(answer_verticals[0], 3), "75")

            self.assertEqual(self._row_text(question_verticals[1], 2), "−38")
            self.assertEqual(self._row_text(answer_verticals[1], 0), "借")
            self.assertEqual(self._row_text(answer_verticals[1], 3), "14")

            bottom_border = (
                question_verticals[0]
                .cell(2, 2)
                ._tc.get_or_add_tcPr()
                .first_child_found_in("w:tcBorders")
                .find(qn("w:bottom"))
            )
            self.assertIsNotNone(bottom_border)
            self.assertEqual(bottom_border.get(qn("w:val")), "single")
            self.assertEqual(bottom_border.get(qn("w:sz")), "8")
            paragraph_borders = (
                question_verticals[0]
                .cell(2, 2)
                .paragraphs[0]
                ._p.get_or_add_pPr()
                .find(qn("w:pBdr"))
            )
            self.assertIsNone(paragraph_borders)

    def test_operator_stays_next_to_widest_operand_with_matching_style(self) -> None:
        problems = iter(
            [
                VerticalArithmeticProblem(
                    operation="addition",
                    left=47,
                    right=58,
                    result=105,
                    working=AdditionWorking((1, 1, None)),
                ),
                VerticalArithmeticProblem(
                    operation="subtraction",
                    left=152,
                    right=38,
                    result=114,
                    working=SubtractionWorking((False, False, True)),
                ),
            ]
        )
        builder = VerticalArithmeticDocumentBuilder(
            {
                "pages": 1,
                "count": 2,
                "columns": 1,
                "title": "运算符定位测试",
                "output_file": "operators.docx",
                "output_file_answer": "operators-answer.docx",
                "font_name": "黑体",
                "font_size": 18,
                "show_working_in_answer": True,
                "hard_label": False,
            },
            asset_dir=PROJECT_ROOT / "src",
            rng=random.Random(5),
        )

        with tempfile.TemporaryDirectory() as temp_dir:
            pair = builder.build(lambda: next(problems), output_dir=temp_dir)
            document = Document(pair.question)
            outer_table = document.tables[0]
            addition_table = outer_table.cell(0, 0).tables[0]
            subtraction_table = outer_table.cell(1, 0).tables[0]

            addition_operator = addition_table.cell(2, 2)
            subtraction_operator = subtraction_table.cell(2, 1)
            self.assertEqual(addition_operator.text.strip(), "+")
            self.assertEqual(subtraction_operator.text.strip(), "−")

            for operator_cell in (addition_operator, subtraction_operator):
                operator_run = operator_cell.paragraphs[0].runs[0]
                self.assertEqual(operator_run.font.name, "Arial")
                self.assertEqual(operator_run.font.size.pt, 18.0)

            addition_leading_cell = addition_table.cell(2, 1)
            self.assertEqual(addition_leading_cell.text, "")
            self.assertIsNone(
                addition_leading_cell
                ._tc.get_or_add_tcPr()
                .first_child_found_in("w:tcBorders")
            )

            for table, operator_column in (
                (addition_table, 2),
                (subtraction_table, 1),
            ):
                border = (
                    table.cell(2, operator_column)
                    ._tc.get_or_add_tcPr()
                    .first_child_found_in("w:tcBorders")
                    .find(qn("w:bottom"))
                )
                self.assertEqual(border.get(qn("w:val")), "single")

    @staticmethod
    def _row_text(table, row_index: int) -> str:
        return "".join(cell.text.strip() for cell in table.rows[row_index].cells)


class VerticalArithmeticFlowTests(unittest.TestCase):
    def test_flow_builds_docx_without_pdf_when_conversion_is_disabled(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            config = {
                "app": {
                    "output_dir": temp_dir,
                    "convert_to_pdf": False,
                    "random_seed": 23,
                },
                "flows": {
                    "vertical_arithmetic": {
                        "pages": 1,
                        "count": 4,
                        "columns": 2,
                        "title": "竖式工作流测试",
                        "output_file": "questions.docx",
                        "output_file_answer": "answers.docx",
                        "hard_label": False,
                        "settings": [
                            {
                                "operation": "addition",
                                "left_min": 10,
                                "left_max": 99,
                                "right_min": 10,
                                "right_max": 99,
                                "carry": "required",
                            },
                            {
                                "operation": "subtraction",
                                "left_min": 10,
                                "left_max": 99,
                                "right_min": 10,
                                "right_max": 99,
                                "borrow": "required",
                            },
                        ],
                    }
                },
            }
            context = AppContext(
                PROJECT_ROOT,
                config,
                "vertical_arithmetic",
                logging.getLogger("vertical-flow-test"),
            )

            result = run(context)

            self.assertEqual(len(result.docx_files), 2)
            self.assertEqual(result.pdf_files, ())
            self.assertTrue(all(path.is_file() for path in result.docx_files))


if __name__ == "__main__":
    unittest.main()
